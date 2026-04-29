import os
from PIL import Image, ImageDraw, ImageFont
from fpdf import FPDF
from docx import Document
from openpyxl import Workbook
from gtts import gTTS

# Create directory
OUTPUT_DIR = "/Users/macbook/Documents/visatamora/tests/sample_files"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Common font path on Mac
FONT_PATH = "/System/Library/Fonts/Supplemental/Arial.ttf"

def get_font(size):
    try:
        return ImageFont.truetype(FONT_PATH, size)
    except:
        return ImageFont.load_default()

def generate_image():
    # 1. struk_restoran.jpg
    img = Image.new('RGB', (600, 800), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    font_header = get_font(40)
    font_body = get_font(30)
    
    d.text((50, 50), "WARUNG NASI GORENG BAROKAH", fill=(0, 0, 0), font=font_header)
    d.text((50, 110), "=========================", fill=(0, 0, 0), font=font_body)
    d.text((50, 160), "Nasi Goreng Spesial  Rp 25.000", fill=(0, 0, 0), font=font_body)
    d.text((50, 210), "Es Teh Manis         Rp  5.000", fill=(0, 0, 0), font=font_body)
    d.text((50, 260), "=========================", fill=(0, 0, 0), font=font_body)
    d.text((50, 310), "TOTAL                Rp 30.000", fill=(0, 0, 0), font=font_body)
    d.text((50, 370), "Bayar: QRIS", fill=(0, 0, 0), font=font_body)
    d.text((50, 420), "Tanggal: 2026-04-29", fill=(0, 0, 0), font=font_body)
    
    img.save(os.path.join(OUTPUT_DIR, "struk_restoran.jpg"))
    print("Generated struk_restoran.jpg with custom font")

def generate_pdf_text():
    # 2. invoice_indihome.pdf
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    text = [
        "INVOICE - INV/2026/04/001",
        "Dari: PT Telkom Indonesia (IndiHome)",
        "Kepada: Budi Santoso",
        "Layanan: Internet Fiber 50Mbps",
        "Periode: April 2026",
        "Total: Rp 250.000",
        "Jatuh Tempo: 2026-05-10"
    ]
    for line in text:
        pdf.cell(200, 10, text=line, new_x="LMARGIN", new_y="NEXT")
    pdf.output(os.path.join(OUTPUT_DIR, "invoice_indihome.pdf"))
    print("Generated invoice_indihome.pdf")

def generate_docx():
    # 3. belanja_april.docx
    doc = Document()
    doc.add_heading('Laporan Belanja - April 2026', 0)
    doc.add_paragraph('Merchant: Indomaret')
    doc.add_paragraph('Tanggal: 2026-04-29')
    doc.add_paragraph('Item:')
    doc.add_paragraph('- Minyak Goreng 2L   Rp 28.000')
    doc.add_paragraph('- Beras 5kg          Rp 75.000')
    doc.add_paragraph('- Sabun Mandi 3pcs   Rp 18.000')
    doc.add_paragraph('Total: Rp 121.000')
    doc.add_paragraph('Bayar: Debit BCA')
    doc.save(os.path.join(OUTPUT_DIR, "belanja_april.docx"))
    print("Generated belanja_april.docx")

def generate_xlsx():
    # 4. keanggotaan_gym.xlsx
    wb = Workbook()
    ws = wb.active
    ws.append(['Merchant', 'Tanggal', 'Keterangan', 'Jumlah'])
    ws.append(['FitLife Gym', '2026-04-29', 'Membership Bulanan', 350000])
    ws.append(['', '', 'Bayar: Transfer BCA', ''])
    wb.save(os.path.join(OUTPUT_DIR, "keanggotaan_gym.xlsx"))
    print("Generated keanggotaan_gym.xlsx")

def generate_audio():
    # 5. catatan_belanja.mp3
    text = "Hari ini isi bensin Pertamax 20 liter total 420 ribu rupiah, bayar pakai kartu debit di SPBU Pertamina Jalan Sudirman."
    tts = gTTS(text=text, lang='id')
    tts.save(os.path.join(OUTPUT_DIR, "catatan_belanja.mp3"))
    print("Generated catatan_belanja.mp3")

def generate_pdf_scan():
    # 6. kwitansi_scan.pdf
    # First create an image of the receipt
    img = Image.new('RGB', (800, 400), color=(255, 255, 240))
    d = ImageDraw.Draw(img)
    font_header = get_font(40)
    font_body = get_font(25)
    
    d.text((50, 40), "KWITANSI", fill=(0, 0, 128), font=font_header)
    d.text((50, 100), "No: 001/IV/2026", fill=(0, 0, 128), font=font_body)
    d.text((50, 140), "Sudah terima dari: Budi Santoso", fill=(0, 0, 128), font=font_body)
    d.text((50, 180), "Uang sejumlah: Tiga Ratus Ribu Rupiah", fill=(0, 0, 128), font=font_body)
    d.text((50, 220), "Untuk pembayaran: Les Privat Matematika - April 2026", fill=(0, 0, 128), font=font_body)
    d.text((50, 260), "Terbilang: Rp 300.000,-", fill=(0, 0, 128), font=font_body)
    d.text((50, 310), "Jakarta, 2026-04-29", fill=(0, 0, 128), font=font_body)
    
    img_path = os.path.join(OUTPUT_DIR, "kwitansi_temp.jpg")
    img.save(img_path)
    
    # Then put it in a PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.image(img_path, x=10, y=10, w=180)
    pdf.output(os.path.join(OUTPUT_DIR, "kwitansi_scan.pdf"))
    os.remove(img_path)
    print("Generated kwitansi_scan.pdf with custom font")

if __name__ == "__main__":
    generate_image()
    generate_pdf_text()
    generate_docx()
    generate_xlsx()
    generate_audio()
    generate_pdf_scan()
    print("\nAll sample files generated in:", OUTPUT_DIR)
